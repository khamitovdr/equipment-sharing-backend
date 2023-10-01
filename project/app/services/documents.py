import hashlib

from docx import Document
from docx.enum.text import WD_COLOR_INDEX

from app.models.orders import Order


MONTHS = {
    1: "января",
    2: "февраля",
    3: "марта",
    4: "апреля",
    5: "мая",
    6: "июня",
    7: "июля",
    8: "августа",
    9: "сентября",
    10: "октября",
    11: "ноября",
    12: "декабря",
}

PLUG = "_____________"

CONTRACT_MANUFACTURING_TEMPLATES = {
    "Спецтехника": "app/data/contract_templates/special_equipment_template.docx",
    "Промышленное оборудование": "app/data/contract_templates/general_equipment_template.docx",
    "Контрактное производство": "app/data/contract_templates/contract_manufacturing_template.docx",
    "Выставочное оборудование": "app/data/contract_templates/general_equipment_template.docx",
}


async def get_contract_template(
        city: str,
        order: Order,
) -> str:
    renter = await order.requester
    equipment = await order.equipment
    owner = await equipment.organization
    category = await equipment.category

    start_date = order.start_date
    end_date = order.end_date

    renter_organization = await renter.organization
    renter_has_organization = renter_organization is not None

    if renter_has_organization:
        renter_requisites = await renter_organization.requisites
    else:
        renter_requisites = await renter.requisites
    renter_has_requisites = renter_requisites is not None

    owner_requisites = await owner.requisites
    owner_has_requisites = owner_requisites is not None

    substitutions = {
        "order_cost": f"{order.cost:.2f}",
        "equipment_name": equipment.name,
        "city": city,
        "start_day": start_date.day,
        "start_month": MONTHS[start_date.month],
        "start_year": start_date.year,
        "end_day": end_date.day,
        "end_month": MONTHS[end_date.month],
        "end_year": end_date.year,

        "renter_organization_name": renter_organization.short_name if renter_has_organization else PLUG,
        "renter_name": renter.full_name,
        "renter_legal_address": renter_organization.legal_address if renter_has_organization else PLUG,
        "renter_inn": renter_organization.inn if renter_has_organization else PLUG,
        "renter_kpp": renter_organization.kpp if renter_has_organization else PLUG,

        "renter_payment_account": renter_requisites.payment_account if renter_has_requisites else PLUG,
        "renter_bank_name": renter_requisites.bank_name if renter_has_requisites else PLUG,
        "renter_bank_bic": renter_requisites.bank_bic if renter_has_requisites else PLUG,
        "renter_bank_correspondent_account": renter_requisites.bank_correspondent_account if renter_has_requisites else PLUG,

        "owner_organization_name": owner.short_name,
        "owner_manager_name": owner.manager_name,
        "owner_legal_address": owner.legal_address,
        "owner_inn": owner.inn,
        "owner_kpp": owner.kpp,
        
        "owner_payment_account": owner_requisites.payment_account if owner_has_requisites else PLUG,
        "owner_bank_name": owner_requisites.bank_name if owner_has_requisites else PLUG,
        "owner_bank_bic": owner_requisites.bank_bic if owner_has_requisites else PLUG,
        "owner_bank_correspondent_account": owner_requisites.bank_correspondent_account if owner_has_requisites else PLUG,
    }

    contract_template_path = CONTRACT_MANUFACTURING_TEMPLATES[category.name]
    doc = Document(contract_template_path)

    for paragraph in doc.paragraphs:
        for key, value in substitutions.items():
            if f"<<{key}>>" in paragraph.text:
                paragraph.text = paragraph.text.replace(f"<<{key}>>", str(value))
                paragraph.runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in substitutions.items():
                    if f"<<{key}>>" in cell.text:
                        cell.text = cell.text.replace(f"<<{key}>>", str(value))
                        cell.paragraphs[0].runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW

    hash = hashlib.sha1(str(substitutions).encode('utf-8')).hexdigest()
    save_path = f'static/orders/documents/{hash}.docx'
    doc.save(save_path)
    return save_path
